import numpy as np
import matplotlib.pyplot as plt
import sounddevice as sd
import soundfile as sf
import scipy.fftpack
from docx import Document
from docx.shared import Inches
from io import BytesIO

data, fs = sf.read('sound1.wav', dtype='float32')
print(data.dtype)
print(data.shape)
sd.play(data, fs)
status = sd.wait()

# ==============================
# ========== [ZAD1] ============
# ==============================

# ROZDZIELENIE KANAŁÓW NA L, R i MIX
sf.write('sound_L.wav', data[:,0], fs)
sf.write('sound_R.wav', data[:,1], fs)
sf.write('sound_mix.wav', ((data[:,0] + data[:,1]) / 2), fs)

# WYKRESY KANAŁÓW L, R, MIX
plt.figure()
plt.subplot(3,1,1)
plt.title("Kanał L")
plt.plot(data[:,0])
plt.subplot(3,1,2)
plt.title("Kanał R")
plt.plot(data[:,1])
plt.subplot(3,1,3)
plt.title("Kanał MIX")
plt.plot((data[:,0] + data[:,1]) / 2)
plt.show()

# WYKRES CZASU (WIDMO)
data, fs = sf.read('sin_440Hz.wav', dtype=np.int32)
plt.figure()
plt.subplot(2,1,1)
plt.plot(np.arange(0,data.shape[0])/fs,data)
plt.xlim([0, 0.02])

# WYKRES CZĘSTOTLIWOŚCI (WIDMO)
plt.subplot(2,1,2)
yf = scipy.fftpack.fft(data)
plt.plot(np.arange(0,fs,1.0*fs/yf.size),np.abs(yf))
plt.show()

fsize=2**8

# MODYFIKACJA WIDMA
plt.figure()
plt.subplot(2,1,1)
plt.plot(np.arange(0,data.shape[0])/fs,data)
plt.xlim([0, 0.02])
plt.subplot(2,1,2)
yf = scipy.fftpack.fft(data,fsize)
plt.plot(np.arange(0,fs,fs/fsize),np.abs(yf))
plt.show()

# POŁOWA WIDMA
plt.figure()
plt.subplot(2,1,1)
plt.plot(np.arange(0,data.shape[0])/fs,data)
plt.xlim([0, 0.02])
plt.subplot(2,1,2)
yf = scipy.fftpack.fft(data,fsize)
plt.plot(np.arange(0,fs/2,fs/fsize),np.abs(yf[:fsize//2]))
plt.show()

# WIDMO W SKALI DECYBELOWEJ
plt.figure()
plt.subplot(2,1,1)
plt.plot(np.arange(0,data.shape[0])/fs,data)
plt.xlim([0, 0.02])
plt.subplot(2,1,2)
yf = scipy.fftpack.fft(data,fsize)
plt.plot(np.arange(0,fs/2,fs/fsize),20*np.log10( np.abs(yf[:fsize//2])))
plt.show()

# ==============================
# ========== [ZAD2] ============
# ==============================

def plotAudio(Signal, Fs, TimeMargin=[0, 0.02], fsize=2**8, axs=None):

    if axs is None:
        fig, axs = plt.subplots(2, 1, figsize=(10, 7))

    time_axis = np.arange(len(Signal)) / Fs
    axs[0].plot(time_axis, Signal)
    axs[0].set_title("Sygnał w czasie")
    axs[0].set_xlabel("Czas [s]")
    axs[0].set_ylabel("Amplituda")
    axs[0].set_xlim(TimeMargin)
    axs[0].grid(True)

    yf = scipy.fftpack.fft(Signal, fsize)
    yf_polowa = np.abs(yf[:fsize // 2])
    yf_decybele = 20 * np.log10(yf_polowa)

    freq_axis = np.linspace(0, Fs / 2, fsize // 2)

    axs[1].plot(freq_axis, yf_decybele)
    axs[1].set_title(f"Połowa widma [fsize: {fsize}]")
    axs[1].set_xlabel("Częstotliwość [Hz]")
    axs[1].set_ylabel("Magnituda [dB]")
    axs[1].grid(True)

    max_idx = np.argmax(yf_polowa)
    naj_czestotliwosc = freq_axis[max_idx]
    naj_wartosc_amp = yf_decybele[max_idx]

    return naj_czestotliwosc, naj_wartosc_amp

data_440, fs_440 = sf.read('sin_440Hz.wav')
plotAudio(data_440, fs_440)
plt.show()

# ==============================
# ========== [ZAD3] ============
# ==============================

document = Document()
document.add_heading('Julian Vrtiska Sprawozdanie LAB1 - Zadanie 3', 0)

pliki = ['sin_60Hz.wav', 'sin_440Hz.wav', 'sin_8000Hz.wav']
fsizes = [2**8, 2**12, 2**16]
TimeMargin = [0, 0.02]

for plik in pliki:

    data, fs = sf.read(plik)
    if len(data.shape) > 1:
        data = data[:, 0]

    document.add_heading(f'Plik - {plik}', 2)

    for fsize in fsizes:
        document.add_heading(f'Rozmiar widma [fsize: {fsize}]', 3)
        fig, axs = plt.subplots(2, 1, figsize=(10, 7))

        # WYWOŁANIE FUNKCJI Z ZADANIA 2
        f_max, a_max = plotAudio(data, fs, TimeMargin=TimeMargin, fsize=fsize, axs=axs)

        # TYTUŁ WYKRESU
        fig.suptitle(f'Plik {plik} [fsize: {fsize}]')
        # POPRAWA CZYTELNOŚCI
        fig.tight_layout(pad=1.5)

        # TWORZENIE BUFORA
        bufor = BytesIO()
        # ZAPIS DO BUFORA
        fig.savefig(bufor)
        # DODANIE OBRAZU Z BUFORA DO PLIKU
        document.add_picture(bufor, width=Inches(6))

        bufor.close()
        plt.close(fig)

        document.add_paragraph(f'Najwyższa wartość widma dla częstotliwości: {f_max:.2f} Hz')
        document.add_paragraph(f'Wartość amplitudy w tym miejscu: {a_max:.4f} dB')

# GENEROWANIE DOCX
docx_path = 'julian_vrtiska_sprawozdanie_lab1.docx'
document.save(docx_path)